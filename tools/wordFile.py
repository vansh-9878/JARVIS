from docx import Document
import google.generativeai as genai
from openai import OpenAI
import os
from langchain_core.tools import tool
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(
    api_key=os.getenv("GEMINI_API23"),
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
)

def generateContent(user_query):
    systemMessage="You are an AI assistant that generates essays and reports on a given topic, the content should be professional and in paragraphs only"
    messages = [{"role": "user", "content": systemMessage},{"role": "user", "content": user_query}]
    response = client.chat.completions.create(
        model="gemini-2.5-flash-lite",
        messages=messages
        )
    return response.choices[0].message.content

@tool
def storeFile(user_query:str,heading:str,filePath:str="C:\\Users\\Niall Dcunha\\Downloads")->str:
    """Generates content based on the user query, stores it in a word file and then opens it"""
    content=generateContent(user_query)
    doc = Document()
    doc.add_heading(heading, 0)
    filepath=os.path.join(filePath,heading+".docx")
    doc.add_paragraph(content)
    doc.save(filepath)
    print(f"Report saved to {filepath}")
    os.startfile(filepath)
    return f"file saved at {filepath}"
